"""
Parse Word doc for Round 2 - extract ALL 4 tables into proper JSON structure
"""
import docx
import json

doc = docx.Document(r'd:\code\Supervision_phut_1-2569-database\ประเด็นการนิเทศงาน ปี 69 จังหวัดนครปฐม รวม.docx')

# ===== TABLE 0: ประเด็นการนิเทศงาน =====
table0 = doc.tables[0]
groups = []
current_group = None

for r_idx, row in enumerate(table0.rows):
    cells = [cell.text.strip() for cell in row.cells]
    
    if r_idx == 0:
        # header row
        continue
    
    # Check if this is a section header (merged row - all cells have same text)
    unique_texts = set(cells)
    if len(unique_texts) == 1 and len(cells[0]) > 10:
        # This is a group header
        current_group = {
            "id": str(len(groups) + 1),
            "title": cells[0],
            "projects": []
        }
        groups.append(current_group)
        continue
    
    # Regular data row
    if current_group is None:
        current_group = {
            "id": "1",
            "title": "ประเด็นการนิเทศงาน",
            "projects": []
        }
        groups.append(current_group)
    
    project = {
        "name": cells[0] if len(cells) > 0 else "",
        "result": cells[1] if len(cells) > 1 else "",
        "progress": cells[2] if len(cells) > 2 else "",
        "problem": cells[3] if len(cells) > 3 else "",
        "solution": cells[4] if len(cells) > 4 else "",
        "images": []
    }
    current_group["projects"].append(project)

# Save Table 0
with open(r'd:\code\Supervision_phut_1-2569-database\src\data\detailedBudgetProjects.round2.json', 'w', encoding='utf-8') as f:
    json.dump(groups, f, ensure_ascii=False, indent=4)
print(f"TABLE 0: Saved {len(groups)} groups")
for g in groups:
    print(f"  Group '{g['id']}': {g['title'][:60]}... ({len(g['projects'])} projects)")

# ===== TABLE 1: งบประมาณและครุภัณฑ์ =====
table1 = doc.tables[1]
budget_items = []

for r_idx, row in enumerate(table1.rows):
    cells = [cell.text.strip() for cell in row.cells]
    if r_idx == 0:
        continue  # header
    
    # Check if merged header row
    unique_texts = set(cells)
    if len(unique_texts) == 1:
        continue  # skip section headers
    
    budget_items.append({
        "topic": cells[0] if len(cells) > 0 else "",
        "result": cells[1] if len(cells) > 1 else "",
        "problem": cells[2] if len(cells) > 2 else ""
    })

# We'll store budget info as a structured object for round 2
budget_round2 = {
    "budgetItems": budget_items
}

with open(r'd:\code\Supervision_phut_1-2569-database\src\data\budgetData.round2.json', 'w', encoding='utf-8') as f:
    json.dump(budget_round2, f, ensure_ascii=False, indent=4)
print(f"\nTABLE 1: Saved {len(budget_items)} budget items")
for b in budget_items:
    print(f"  - {b['topic'][:80]}...")

# ===== TABLE 2: ครุภัณฑ์โครงการงบยุทธศาสตร์ =====
table2 = doc.tables[2]
equipment_items = []

for r_idx, row in enumerate(table2.rows):
    cells = [cell.text.strip() for cell in row.cells]
    if r_idx == 0:
        continue  # header
    
    equipment_items.append({
        "id": int(cells[0]) if cells[0].isdigit() else r_idx,
        "name": cells[1] if len(cells) > 1 else "",
        "status": cells[2] if len(cells) > 2 else "",
        "statusText": cells[2] if len(cells) > 2 else "",
        "details": cells[3] if len(cells) > 3 else "",
        "problem": "-",
        "solution": "-"
    })

with open(r'd:\code\Supervision_phut_1-2569-database\src\data\projectAssets.round2.json', 'w', encoding='utf-8') as f:
    json.dump(equipment_items, f, ensure_ascii=False, indent=4)
print(f"\nTABLE 2: Saved {len(equipment_items)} equipment items")

# ===== TABLE 3: เรื่องอื่นๆ =====
table3 = doc.tables[3]
other_issues = []
for r_idx, row in enumerate(table3.rows):
    cells = [cell.text.strip() for cell in row.cells]
    if r_idx == 0:
        continue
    other_issues.append({
        "id": str(r_idx),
        "title": "เรื่องอื่นๆ",
        "content": cells[0] if len(cells) > 0 else ""
    })

with open(r'd:\code\Supervision_phut_1-2569-database\src\data\otherIssues.round2.json', 'w', encoding='utf-8') as f:
    json.dump(other_issues, f, ensure_ascii=False, indent=4)
print(f"\nTABLE 3: Saved {len(other_issues)} other issues")

print("\n✅ Done! All Round 2 data extracted from Word document.")
