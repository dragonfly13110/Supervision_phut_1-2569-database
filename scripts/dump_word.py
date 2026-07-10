import docx
import json

doc = docx.Document(r'd:\code\Supervision_phut_1-2569-database\ประเด็นการนิเทศงาน ปี 69 จังหวัดนครปฐม รวม.docx')

print(f'=== TOTAL TABLES: {len(doc.tables)} ===\n')

# Print all paragraphs to understand document structure
print('=== PARAGRAPHS (first 50) ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f'  P{i}: [{para.style.name}] {text[:120]}')
    if i > 100:
        break

print('\n\n')

for t_idx, table in enumerate(doc.tables):
    print(f'\n{"="*80}')
    print(f'=== TABLE {t_idx} === Rows: {len(table.rows)}, Cols: {len(table.columns)}')
    print(f'{"="*80}')
    for r_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            text = cell.text.strip().replace('\n', '\\n')
            cells.append(text[:100])
        print(f'  Row {r_idx}: {cells}')
        if r_idx > 30:
            print(f'  ... ({len(table.rows) - r_idx - 1} more rows)')
            break
